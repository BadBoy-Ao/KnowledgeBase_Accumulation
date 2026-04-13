from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
import subprocess
from uuid import uuid4
import zipfile

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from PIL import Image, ImageFont
import requests

from app.config import (
    DEFAULT_FONT_CANDIDATES,
    DOCX_DIR,
    IMAGES_DIR,
    JAVA_REPORT_TEMPLATE_ROOT,
    PDF_DIR,
    SEAL_API_BASE_URL,
    SEALS_DIR,
    ZIPS_DIR,
)
from app.models.report import BatchReportRequest, ReportEntry
from app.services.template_registry import (
    TemplateConfig,
    build_beijing_registry,
    build_shandong_registry,
)


PAGE_WIDTH = 1240
PAGE_HEIGHT = 1754
BODY_FONT_SIZE = 34
FOOTER_RIGHT_MARGIN = 182
STAMP_BOX_X = 510
STAMP_BOX_WIDTH = 488
STAMP_BOX_HEIGHT = 235


@dataclass
class BatchRenderResult:
    task_name: str
    template_type: str
    file_name: str
    zip_path: Path
    image_paths: list[Path]
    docx_paths: list[Path]
    pdf_paths: list[Path]


class ReportRenderService:
    def __init__(self) -> None:
        self.font_path = self._resolve_font_path()
        self.body_font = self._load_font(BODY_FONT_SIZE)
        self.registry = {
            **build_shandong_registry(JAVA_REPORT_TEMPLATE_ROOT),
            **build_beijing_registry(JAVA_REPORT_TEMPLATE_ROOT),
        }

    def generate_zip(self, request: BatchReportRequest) -> BatchRenderResult:
        config = self.get_template_config(request.templateType)
        self._validate_request(config, request)
        normalized_shared = self._build_shared_payload(config, request.sharedPayload)

        docx_paths: list[Path] = []
        pdf_paths: list[Path] = []
        image_paths: list[Path] = []

        items = self._expand_generation_items(
            config, request.entries, normalized_shared
        )
        for index, item in enumerate(items, start=1):
            docx_path = self._render_single_docx(
                config, item["render_data"], item["table_rows"], index
            )
            pdf_path = self._convert_docx_to_pdf(docx_path)
            image_path = self._convert_pdf_to_image(
                pdf_path=pdf_path,
                index=index,
                config=config,
                render_data=item["render_data"],
                file_label=item["file_label"],
                primary_seal_company=item.get("primary_seal_company"),
                secondary_seal_company=item.get("secondary_seal_company"),
                primary_footer_company=item.get("primary_footer_company"),
                secondary_footer_company=item.get("secondary_footer_company"),
            )
            docx_paths.append(docx_path)
            pdf_paths.append(pdf_path)
            image_paths.append(image_path)

        file_name = f"{config.zip_prefix}_{uuid4().hex[:12]}.zip"
        zip_path = self._pack_zip(file_name, image_paths)
        return BatchRenderResult(
            task_name=self._build_task_name(config, request.entries),
            template_type=config.template_type,
            file_name=file_name,
            zip_path=zip_path,
            image_paths=image_paths,
            docx_paths=docx_paths,
            pdf_paths=pdf_paths,
        )

    def get_template_config(self, template_type: str) -> TemplateConfig:
        if template_type not in self.registry:
            raise ValueError(f"不支持的模板类型: {template_type}")
        return self.registry[template_type]

    def _validate_request(
        self, config: TemplateConfig, request: BatchReportRequest
    ) -> None:
        if not request.entries:
            raise ValueError("entries 不能为空")

        for field_name in config.required_shared_fields:
            value = request.sharedPayload.get(field_name)
            if (
                value is None
                or (isinstance(value, str) and not value.strip())
                or (isinstance(value, list) and not value)
            ):
                raise ValueError(f"sharedPayload.{field_name} 不能为空")

        for index, entry in enumerate(request.entries, start=1):
            if not entry.companyName.strip():
                raise ValueError(f"第{index}条 entry.companyName 不能为空")

        if config.template_type in {
            "shandong_short_link_reporting",
            "shandong_multi_level_domain",
            "shandong_phone_ownership_claim",
            "beijing_number_ownership_claim",
        }:
            for index, entry in enumerate(request.entries, start=1):
                if not (entry.signature or "").strip():
                    raise ValueError(f"第{index}条 entry.signature 不能为空")

    def _build_shared_payload(
        self, config: TemplateConfig, shared_payload: dict[str, object]
    ) -> dict[str, object]:
        merged: dict[str, object] = dict(shared_payload)
        for key, value in config.default_shared_payload.items():
            if key not in merged or merged[key] in (None, "", []):
                merged[key] = value() if callable(value) else value
        return merged

    def _expand_generation_items(
        self,
        config: TemplateConfig,
        entries: list[ReportEntry],
        shared_payload: dict[str, object],
    ) -> list[dict[str, object]]:
        items: list[dict[str, object]] = []
        if config.template_type == "beijing_compliance_commitment":
            rows = self._build_beijing_compliance_rows(shared_payload.get("lgiItems"))
            if not rows:
                raise ValueError("sharedPayload.lgiItems 解析后为空")
            chunks = [rows[i : i + 10] for i in range(0, len(rows), 10)]
            for entry in entries:
                for chunk_index, chunk_rows in enumerate(chunks, start=1):
                    render_data = {
                        "companyName": entry.companyName.strip(),
                        "businessScope": str(
                            shared_payload.get("businessScope", "金融业务")
                        ).strip(),
                        "time": str(shared_payload.get("time", "")).strip(),
                    }
                    items.append(
                        {
                            "render_data": render_data,
                            "table_rows": chunk_rows,
                            "file_label": f"{entry.companyName.strip()}_{chunk_index}",
                            "primary_seal_company": entry.companyName.strip(),
                            "secondary_seal_company": str(
                                shared_payload.get(
                                    "ourCompanyName", "深圳硕软技术有限公司"
                                )
                            ).strip(),
                            "primary_footer_company": entry.companyName.strip(),
                            "secondary_footer_company": "代理渠道单位",
                        }
                    )
            return items

        if config.template_type == "beijing_number_ownership_claim":
            rows = self._build_beijing_number_rows(shared_payload.get("numberList"))
            if not rows:
                raise ValueError("sharedPayload.numberList 解析后为空")
            for entry in entries:
                items.append(
                    {
                        "render_data": {
                            "companyName": entry.companyName.strip(),
                            "signature": (entry.signature or "").strip(),
                            "time": str(shared_payload.get("time", "")).strip(),
                        },
                        "table_rows": rows,
                        "file_label": (entry.signature or entry.companyName).strip(),
                        "primary_seal_company": entry.companyName.strip(),
                        "primary_footer_company": entry.companyName.strip(),
                    }
                )
            return items

        for entry in entries:
            render_data = self._build_render_data(config, entry, shared_payload)
            item = {
                "render_data": render_data,
                "table_rows": None,
                "file_label": (entry.signature or entry.companyName).strip(),
                "primary_seal_company": None,
                "secondary_seal_company": None,
                "primary_footer_company": None,
                "secondary_footer_company": None,
            }

            if config.template_type.startswith("shandong_"):
                item["primary_seal_company"] = str(
                    render_data.get(config.seal_company_field, "")
                ).strip()
                item["primary_footer_company"] = str(
                    render_data.get(config.seal_company_field, "")
                ).strip()
            elif config.template_type == "beijing_link_authorization":
                item["primary_seal_company"] = str(
                    render_data.get("linkCompanyName", "")
                ).strip()
                item["secondary_seal_company"] = str(
                    render_data.get("signatureCompanyName", "")
                ).strip()
                item["primary_footer_company"] = str(
                    render_data.get("linkCompanyName", "")
                ).strip()
                item["secondary_footer_company"] = str(
                    render_data.get("signatureCompanyName", "")
                ).strip()

            items.append(item)
        return items

    def _build_render_data(
        self,
        config: TemplateConfig,
        entry: ReportEntry,
        shared_payload: dict[str, object],
    ) -> dict[str, str]:
        data: dict[str, str] = {}
        for key, value in shared_payload.items():
            multiline = key in config.multiline_fields
            data[key] = self._normalize_value(value, multiline=multiline)

        if config.template_type == "shandong_phone_ownership_claim":
            data["phoneNumbers"] = self._normalize_joined_list(
                shared_payload.get("phoneNumbers"), ","
            )
        if config.template_type == "beijing_link_authorization":
            data["linkList"] = self._normalize_joined_list(
                shared_payload.get("linkList"), "、"
            )

        data["companyName"] = entry.companyName.strip()
        if entry.signature:
            data["signature"] = entry.signature.strip()

        if config.template_type == "shandong_link_authorization":
            data["signatureCompanyName"] = entry.companyName.strip()
        if config.template_type == "beijing_link_authorization":
            data["signatureCompanyName"] = entry.companyName.strip()
        return data

    def _render_single_docx(
        self,
        config: TemplateConfig,
        render_data: dict[str, str],
        table_rows: list[list[str]] | None,
        index: int,
    ) -> Path:
        document = Document(str(config.template_path))
        self._replace_text_placeholders(document, render_data)
        if (
            config.template_type == "beijing_compliance_commitment"
            and table_rows is not None
        ):
            self._insert_table_at_marker(
                document,
                "{{#lgiList}}",
                ["链接/号码", "说明"],
                table_rows,
                font_size=10,
            )
        if (
            config.template_type == "beijing_number_ownership_claim"
            and table_rows is not None
        ):
            self._insert_table_at_marker(
                document, "{{#numberList}}", ["号码", "说明"], table_rows, font_size=14
            )
        docx_path = DOCX_DIR / f"{config.template_type}_{index}_{uuid4().hex[:12]}.docx"
        document.save(str(docx_path))
        self._clear_image_placeholders(docx_path)
        return docx_path

    def _clear_image_placeholders(self, docx_path: Path) -> None:
        with zipfile.ZipFile(docx_path, "r") as source_zip:
            file_map = {
                entry.filename: source_zip.read(entry.filename)
                for entry in source_zip.infolist()
            }

        document_xml = file_map["word/document.xml"].decode("utf-8")
        document_xml = re.sub(r"\{\{@[^}]+\}\}", "", document_xml)
        file_map["word/document.xml"] = document_xml.encode("utf-8")

        with zipfile.ZipFile(
            docx_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as target_zip:
            for file_name, data in file_map.items():
                target_zip.writestr(file_name, data)

    def _replace_text_placeholders(
        self, document: Document, render_data: dict[str, str]
    ) -> None:
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, render_data)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, render_data)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, render_data)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, render_data)

    def _replace_in_paragraph(self, paragraph, render_data: dict[str, str]) -> None:
        if not paragraph.runs:
            return

        joined_text = "".join(run.text for run in paragraph.runs)
        if "{{@" in joined_text and "}}" in joined_text:
            for run in paragraph.runs:
                run.text = ""
            return

        for run in paragraph.runs:
            for key, value in render_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder not in run.text:
                    continue
                if "\n" in value:
                    lines = value.splitlines()
                    run.text = lines[0] if lines else ""
                    for extra_line in lines[1:]:
                        run.add_break()
                        run.add_text(extra_line)
                else:
                    run.text = run.text.replace(placeholder, value)

    def _insert_table_at_marker(
        self,
        document: Document,
        marker: str,
        headers: list[str],
        rows: list[list[str]],
        font_size: int,
    ) -> None:
        marker_paragraph = None
        for paragraph in document.paragraphs:
            if marker in paragraph.text:
                marker_paragraph = paragraph
                break
        if marker_paragraph is None:
            raise ValueError(f"未找到循环标记: {marker}")

        table = document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            self._style_cell_paragraphs(header_cells[idx], font_size)

        for row_values in rows:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row_values):
                row_cells[idx].text = value
                self._style_cell_paragraphs(row_cells[idx], font_size)

        marker_paragraph._p.addnext(table._tbl)
        self._remove_paragraph(marker_paragraph)

    def _style_cell_paragraphs(self, cell, font_size: int) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                if self.font_path is not None:
                    run.font.name = "仿宋_GB2312"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

    def _convert_docx_to_pdf(self, docx_path: Path) -> Path:
        pdf_path = PDF_DIR / f"{docx_path.stem}.pdf"
        quoted_docx = str(docx_path).replace("'", "''")
        quoted_pdf = str(pdf_path).replace("'", "''")
        script = (
            "$word = $null; "
            "try { "
            "$word = New-Object -ComObject Word.Application; "
            "$word.Visible = $false; "
            "$word.DisplayAlerts = 0; "
            f"$doc = $word.Documents.Open('{quoted_docx}'); "
            f"$doc.SaveAs([ref]'{quoted_pdf}', [ref]17); "
            "$doc.Close(); "
            "} finally { if ($word -ne $null) { try { $word.Quit() } catch { } } }"
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if not pdf_path.exists():
            raise RuntimeError(f"Word 转 PDF 失败: {docx_path}")
        return pdf_path

    def _convert_pdf_to_image(
        self,
        pdf_path: Path,
        index: int,
        config: TemplateConfig,
        render_data: dict[str, str],
        file_label: str,
        primary_seal_company: str | None,
        secondary_seal_company: str | None,
        primary_footer_company: str | None,
        secondary_footer_company: str | None,
    ) -> Path:
        document = fitz.open(str(pdf_path))
        page = document.load_page(0)
        page_width = page.rect.width
        page_height = page.rect.height
        pixmap = page.get_pixmap(dpi=150, alpha=False)
        image = Image.frombytes(
            "RGB", [pixmap.width, pixmap.height], pixmap.samples
        ).convert("RGBA")

        dynamic_center = self._locate_dynamic_seal_center(
            page=page,
            template_type=config.template_type,
            render_data=render_data,
            primary_footer_company=primary_footer_company,
        )
        link_authorization_centers = self._locate_link_authorization_seal_centers(
            page=page,
            template_type=config.template_type,
            link_company_name=render_data.get("linkCompanyName", "").strip(),
            signature_company_name=render_data.get("signatureCompanyName", "").strip(),
        )
        compliance_centers = self._locate_compliance_seal_centers(
            page=page,
            template_type=config.template_type,
            time_text=render_data.get("time", "").strip(),
        )
        document.close()

        if config.template_type.startswith("shandong_") and primary_seal_company:
            seal_bytes = self._download_seal_bytes(primary_seal_company)
            if seal_bytes:
                x, y, w, dy = self._estimate_shandong_footer_position(
                    image, primary_footer_company or primary_seal_company
                )
                self._paste_seal_image(
                    image,
                    seal_bytes,
                    x,
                    y,
                    w,
                    dy,
                    center_override=dynamic_center,
                    page_width=page_width,
                    page_height=page_height,
                )

        if config.template_type == "beijing_link_authorization":
            if primary_seal_company:
                seal_bytes = self._download_seal_bytes(primary_seal_company)
                if seal_bytes:
                    primary_center = (
                        link_authorization_centers.get("primary")
                        if link_authorization_centers
                        else None
                    )
                    self._paste_seal_at_ratio(
                        image,
                        seal_bytes,
                        primary_center[0] if primary_center else 0.33,
                        primary_center[1] if primary_center else 0.58,
                        0.17,
                        0.17,
                        0.72,
                        absolute_center=bool(primary_center),
                        page_width=page_width if primary_center else None,
                        page_height=page_height if primary_center else None,
                    )
            if secondary_seal_company:
                seal_bytes = self._download_seal_bytes(secondary_seal_company)
                if seal_bytes:
                    secondary_center = (
                        link_authorization_centers.get("secondary")
                        if link_authorization_centers
                        else None
                    )
                    self._paste_seal_at_ratio(
                        image,
                        seal_bytes,
                        secondary_center[0] if secondary_center else 0.74,
                        secondary_center[1] if secondary_center else 0.82,
                        0.17,
                        0.17,
                        0.72,
                        absolute_center=bool(secondary_center),
                        page_width=page_width if secondary_center else None,
                        page_height=page_height if secondary_center else None,
                    )

        if config.template_type == "beijing_compliance_commitment":
            if secondary_seal_company:
                seal_bytes = self._download_seal_bytes(secondary_seal_company)
                if seal_bytes:
                    right_center = (
                        compliance_centers.get("right") if compliance_centers else None
                    )
                    self._paste_seal_at_ratio(
                        image,
                        seal_bytes,
                        right_center[0] if right_center else 0.64,
                        right_center[1] if right_center else 0.70,
                        0.16,
                        0.16,
                        0.68,
                        absolute_center=bool(right_center),
                        page_width=page_width if right_center else None,
                        page_height=page_height if right_center else None,
                    )
            if primary_seal_company:
                seal_bytes = self._download_seal_bytes(primary_seal_company)
                if seal_bytes:
                    left_center = (
                        compliance_centers.get("left") if compliance_centers else None
                    )
                    self._paste_seal_at_ratio(
                        image,
                        seal_bytes,
                        left_center[0] if left_center else 0.33,
                        left_center[1] if left_center else 0.70,
                        0.16,
                        0.16,
                        0.68,
                        absolute_center=bool(left_center),
                        page_width=page_width if left_center else None,
                        page_height=page_height if left_center else None,
                    )

        if (
            config.template_type == "beijing_number_ownership_claim"
            and primary_seal_company
        ):
            seal_bytes = self._download_seal_bytes(primary_seal_company)
            if seal_bytes:
                center_override = dynamic_center or None
                self._paste_seal_at_ratio(
                    image,
                    seal_bytes,
                    center_override[0] if center_override else 0.63,
                    center_override[1] if center_override else 0.79,
                    0.18,
                    0.18,
                    0.7,
                    absolute_center=bool(center_override),
                    page_width=page_width if center_override else None,
                    page_height=page_height if center_override else None,
                )

        image_path = (
            IMAGES_DIR
            / f"{config.file_prefix}_{self._safe_name(file_label)}_{index}_{uuid4().hex[:8]}.jpg"
        )
        image.convert("RGB").save(image_path, format="JPEG", quality=92)
        return image_path

    def _download_seal_bytes(self, company_name: str) -> bytes | None:
        company_name = (company_name or "").strip()
        if not company_name:
            return None
        cache_path = SEALS_DIR / f"seal_{self._safe_name(company_name)}.png"
        if cache_path.exists():
            return cache_path.read_bytes()
        response = requests.post(
            f"{SEAL_API_BASE_URL.rstrip('/')}/api/seal/download",
            json={"names": [company_name]},
            headers={"Content-Type": "application/json"},
            timeout=120,
        )
        response.raise_for_status()
        seal_bytes = self._extract_first_image_bytes_from_zip(response.content)
        if seal_bytes:
            cache_path.write_bytes(seal_bytes)
        return seal_bytes

    def _extract_first_image_bytes_from_zip(self, zip_bytes: bytes) -> bytes | None:
        with zipfile.ZipFile(BytesIO(zip_bytes), "r") as zip_file:
            for name in zip_file.namelist():
                if not name.endswith("/"):
                    return zip_file.read(name)
        return None

    def _locate_compliance_seal_centers(
        self, page: fitz.Page, template_type: str, time_text: str
    ) -> dict[str, tuple[float, float]]:
        if template_type != "beijing_compliance_commitment":
            return {}

        left_center = self._locate_center_from_label_and_date(
            page,
            label_text="客户单位名称（盖章）",
            time_text=time_text,
            side="left",
        )
        right_center = self._locate_center_from_label_and_date(
            page,
            label_text="代理渠道单位名称（盖章）",
            time_text=time_text,
            side="right",
        )
        result: dict[str, tuple[float, float]] = {}
        if left_center:
            result["left"] = left_center
        if right_center:
            result["right"] = right_center
        return result

    def _locate_link_authorization_seal_centers(
        self,
        page: fitz.Page,
        template_type: str,
        link_company_name: str,
        signature_company_name: str,
    ) -> dict[str, tuple[float, float]]:
        if template_type != "beijing_link_authorization":
            return {}

        result: dict[str, tuple[float, float]] = {}
        primary_center = self._locate_center_from_company_and_handler(
            page, f"授权方：{link_company_name}", y_anchor="company"
        )
        secondary_center = self._locate_center_from_company_and_handler(
            page, f"被授权方：{signature_company_name}", y_anchor="handler"
        )
        if primary_center:
            result["primary"] = primary_center
        if secondary_center:
            result["secondary"] = secondary_center
        return result

    def _locate_center_from_company_and_handler(
        self, page: fitz.Page, company_line_text: str, y_anchor: str
    ) -> tuple[float, float] | None:
        company_rect = self._find_line_rect(page, company_line_text, "lower")
        if company_rect is None:
            return None

        handler_rect = self._find_next_line_rect(
            page=page,
            anchor_rect=company_rect,
            text="经办人：",
            region="lower",
        )
        if handler_rect is None:
            return (
                (company_rect.x0 + company_rect.x1) / 2,
                company_rect.y1 + (company_rect.y1 - company_rect.y0) * 1.5,
            )

        union_rect = fitz.Rect(
            min(company_rect.x0, handler_rect.x0),
            min(company_rect.y0, handler_rect.y0),
            max(company_rect.x1, handler_rect.x1),
            max(company_rect.y1, handler_rect.y1),
        )
        if y_anchor == "company":
            center_y = (company_rect.y0 + company_rect.y1) / 2
        elif y_anchor == "handler":
            center_y = (handler_rect.y0 + handler_rect.y1) / 2
        else:
            center_y = (union_rect.y0 + union_rect.y1) / 2
        return (
            (union_rect.x0 + union_rect.x1) / 2,
            center_y,
        )

    def _locate_center_from_label_and_date(
        self,
        page: fitz.Page,
        label_text: str,
        time_text: str,
        side: str,
    ) -> tuple[float, float] | None:
        label_rect = self._find_side_line_rect(page, label_text, side)
        date_rect = self._find_side_line_rect(page, f"日期：{time_text}", side)
        if date_rect is None:
            date_rect = self._find_side_line_rect(page, time_text, side)

        if label_rect and date_rect:
            union_rect = fitz.Rect(
                min(label_rect.x0, date_rect.x0),
                min(label_rect.y0, date_rect.y0),
                max(label_rect.x1, date_rect.x1),
                max(label_rect.y1, date_rect.y1),
            )
            return (
                (union_rect.x0 + union_rect.x1) / 2,
                (union_rect.y0 + union_rect.y1) / 2,
            )
        if label_rect:
            return (
                (label_rect.x0 + label_rect.x1) / 2,
                label_rect.y1 + (label_rect.y1 - label_rect.y0) * 1.7,
            )
        return None

    def _estimate_shandong_footer_position(
        self, image: Image.Image, company_name: str
    ) -> tuple[int, int, int, int]:
        scale_x = image.width / PAGE_WIDTH
        scale_y = image.height / PAGE_HEIGHT
        company_prefix = "实际发送主体（盖章）："
        right_x = image.width - int(FOOTER_RIGHT_MARGIN * scale_x)
        font = self._load_font(
            max(20, round(BODY_FONT_SIZE * min(scale_x, scale_y) * 0.78))
        )
        company_line = f"{company_prefix}{company_name}"
        company_line_width = self._text_width(company_line, font)
        prefix_width = self._text_width(company_prefix, font)
        company_width = self._text_width(company_name, font)
        company_x = right_x - company_line_width + prefix_width
        company_y = int(image.height * 0.505)
        date_y = company_y + int(92 * scale_y)
        return company_x, company_y, company_width, date_y

    def _locate_dynamic_seal_center(
        self,
        page: fitz.Page,
        template_type: str,
        render_data: dict[str, str],
        primary_footer_company: str | None,
    ) -> tuple[float, float] | None:
        time_text = render_data.get("time", "").strip()

        if template_type in {
            "shandong_short_link_reporting",
            "shandong_multi_level_domain",
            "shandong_phone_ownership_claim",
            "beijing_number_ownership_claim",
        }:
            company_text = (
                primary_footer_company or render_data.get("companyName", "")
            ).strip()
            return self._locate_center_between_company_and_date(
                page=page,
                company_text=company_text,
                time_text=time_text,
                region="lower",
            )

        if template_type == "shandong_link_authorization":
            company_text = (
                primary_footer_company or render_data.get("linkCompanyName", "")
            ).strip()
            return self._locate_center_between_company_and_date(
                page=page,
                company_text=company_text,
                time_text=time_text,
                region="lower",
            )

        return None

    def _locate_center_between_company_and_date(
        self,
        page: fitz.Page,
        company_text: str,
        time_text: str,
        region: str,
    ) -> tuple[float, float] | None:
        company_rect = self._find_line_rect(page, company_text, region)
        if company_rect is None:
            company_rect = self._find_text_rect(page, company_text, region)

        date_rect = self._find_line_rect(page, time_text, region)
        if date_rect is None:
            date_rect = self._find_text_rect(page, time_text, region)

        if company_rect and date_rect:
            union_rect = fitz.Rect(
                min(company_rect.x0, date_rect.x0),
                min(company_rect.y0, date_rect.y0),
                max(company_rect.x1, date_rect.x1),
                max(company_rect.y1, date_rect.y1),
            )
            return (
                (union_rect.x0 + union_rect.x1) / 2,
                (union_rect.y0 + union_rect.y1) / 2,
            )

        if company_rect:
            return (
                (company_rect.x0 + company_rect.x1) / 2,
                company_rect.y1 + (company_rect.y1 - company_rect.y0) * 0.9,
            )

        return None

    def _find_line_rect(
        self, page: fitz.Page, text: str, region: str
    ) -> fitz.Rect | None:
        normalized_target = self._normalize_match_text(text)
        if not normalized_target:
            return None

        lines = self._extract_line_infos(page, region)
        matches = [
            line for line in lines if normalized_target in line["normalized_text"]
        ]
        if not matches:
            return None
        matches.sort(key=lambda item: (item["rect"].y0, item["rect"].x0))
        return matches[-1]["rect"]

    def _find_side_line_rect(
        self, page: fitz.Page, text: str, side: str
    ) -> fitz.Rect | None:
        normalized_target = self._normalize_match_text(text)
        if not normalized_target:
            return None

        lines = self._extract_line_infos(page, "lower")
        page_mid_x = page.rect.width / 2
        matches = []
        for line in lines:
            rect = line["rect"]
            center_x = (rect.x0 + rect.x1) / 2
            if side == "left" and center_x >= page_mid_x:
                continue
            if side == "right" and center_x < page_mid_x:
                continue
            if normalized_target in line["normalized_text"]:
                matches.append(line)

        if not matches:
            return None
        matches.sort(key=lambda item: (item["rect"].y0, item["rect"].x0))
        return matches[-1]["rect"]

    def _find_next_line_rect(
        self,
        page: fitz.Page,
        anchor_rect: fitz.Rect,
        text: str,
        region: str,
    ) -> fitz.Rect | None:
        normalized_target = self._normalize_match_text(text)
        if not normalized_target:
            return None

        anchor_center_x = (anchor_rect.x0 + anchor_rect.x1) / 2
        matches = []
        for line in self._extract_line_infos(page, region):
            rect = line["rect"]
            if rect.y0 <= anchor_rect.y0:
                continue
            if normalized_target not in line["normalized_text"]:
                continue
            center_x = (rect.x0 + rect.x1) / 2
            if abs(center_x - anchor_center_x) > page.rect.width * 0.22:
                continue
            matches.append(line)

        if not matches:
            return None
        matches.sort(
            key=lambda item: (
                item["rect"].y0,
                abs(((item["rect"].x0 + item["rect"].x1) / 2) - anchor_center_x),
            )
        )
        return matches[0]["rect"]

    def _extract_line_infos(
        self, page: fitz.Page, region: str
    ) -> list[dict[str, object]]:
        words = page.get_text("words")
        lines: dict[tuple[int, int], dict[str, object]] = {}
        for word in words:
            rect = fitz.Rect(word[0], word[1], word[2], word[3])
            if not self._rect_in_region(page, rect, region):
                continue

            key = (word[5], word[6])
            text = str(word[4])
            if key not in lines:
                lines[key] = {"text_parts": [], "rect": rect}
            lines[key]["text_parts"].append(text)
            lines[key]["rect"] |= rect

        result: list[dict[str, object]] = []
        for item in lines.values():
            line_text = "".join(item["text_parts"])
            result.append(
                {
                    "text": line_text,
                    "normalized_text": self._normalize_match_text(line_text),
                    "rect": item["rect"],
                }
            )
        return result

    def _normalize_match_text(self, text: str) -> str:
        return "".join(str(text).split()).strip()

    def _find_text_rect(
        self, page: fitz.Page, text: str, region: str
    ) -> fitz.Rect | None:
        text = (text or "").strip()
        if not text:
            return None

        clip = self._build_search_clip(page, region)
        rects = page.search_for(text, clip=clip)
        if not rects:
            rects = page.search_for(text)
            rects = [rect for rect in rects if self._rect_in_region(page, rect, region)]
        if not rects:
            return None

        rects.sort(key=lambda rect: (rect.y0, rect.x0))
        return rects[-1]

    def _build_search_clip(self, page: fitz.Page, region: str) -> fitz.Rect:
        if region == "lower":
            return fitz.Rect(
                0, page.rect.height * 0.35, page.rect.width, page.rect.height
            )
        return page.rect

    def _rect_in_region(self, page: fitz.Page, rect: fitz.Rect, region: str) -> bool:
        if region == "lower":
            return rect.y0 >= page.rect.height * 0.35
        return True

    def _paste_seal_image(
        self,
        image: Image.Image,
        seal_bytes: bytes,
        company_x: int | float,
        company_y: int | float,
        company_width: int | float,
        date_y: int | float,
        center_override: tuple[float, float] | None = None,
        page_width: float | None = None,
        page_height: float | None = None,
    ) -> None:
        scale_x = image.width / PAGE_WIDTH
        scale_y = image.height / PAGE_HEIGHT
        box_x = int(STAMP_BOX_X * scale_x)
        box_y = int(company_y - 42 * scale_y)
        box_width = int(STAMP_BOX_WIDTH * scale_x)
        box_height = int(STAMP_BOX_HEIGHT * scale_y)
        if center_override and page_width and page_height:
            center_x = int(center_override[0] * image.width / page_width)
            center_y = int(center_override[1] * image.height / page_height)
        else:
            center_x = int(
                max(box_x + box_width * 0.36, company_x + company_width * 0.18)
            )
            center_y = int((company_y + date_y) / 2 + 6 * scale_y)
        self._paste_seal_at_ratio(
            image,
            seal_bytes,
            center_x,
            center_y,
            box_width / image.width,
            box_height / image.height,
            0.68,
            absolute_center=True,
        )

    def _paste_seal_at_ratio(
        self,
        image: Image.Image,
        seal_bytes: bytes,
        x_ratio: float,
        y_ratio: float,
        w_ratio: float,
        h_ratio: float,
        alpha_ratio: float,
        absolute_center: bool = False,
        page_width: float | None = None,
        page_height: float | None = None,
    ) -> None:
        seal_image = Image.open(BytesIO(seal_bytes)).convert("RGBA")
        target_size = int(min(image.width * w_ratio, image.height * h_ratio) * 1.0)
        seal_image.thumbnail((target_size, target_size), Image.Resampling.LANCZOS)
        alpha = seal_image.getchannel("A").point(lambda value: int(value * alpha_ratio))
        seal_image.putalpha(alpha)
        if absolute_center and page_width and page_height:
            center_x = int(x_ratio * image.width / page_width)
            center_y = int(y_ratio * image.height / page_height)
        else:
            center_x = int(x_ratio if absolute_center else image.width * x_ratio)
            center_y = int(y_ratio if absolute_center else image.height * y_ratio)
        paste_x = int(center_x - seal_image.width / 2)
        paste_y = int(center_y - seal_image.height / 2)
        image.alpha_composite(seal_image, (paste_x, paste_y))

    def _build_beijing_compliance_rows(self, items: object) -> list[list[str]]:
        rows: list[list[str]] = []
        if not isinstance(items, list):
            return rows
        for item in items:
            if not isinstance(item, dict):
                continue
            content = str(item.get("content", "")).strip()
            item_type = item.get("type")
            if not content:
                continue
            desc = (
                "我司业务专用链接，申请纳入自名单管理"
                if item_type == 1
                else "我司业务专用号码，申请纳入自名单管理"
            )
            rows.append([content, desc])
        return rows

    def _build_beijing_number_rows(self, value: object) -> list[list[str]]:
        normalized = self._normalize_value(value, multiline=True)
        rows: list[list[str]] = []
        for number in normalized.replace("，", "\n").replace(",", "\n").splitlines():
            number = number.strip()
            if number:
                rows.append([number, "此号码为我司业务专用号码，申请纳入自名单管理"])
        return rows

    def _remove_paragraph(self, paragraph) -> None:
        p = paragraph._element
        p.getparent().remove(p)

    def _pack_zip(self, file_name: str, image_paths: list[Path]) -> Path:
        zip_path = ZIPS_DIR / file_name
        with zipfile.ZipFile(
            zip_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zip_file:
            for image_path in image_paths:
                zip_file.write(image_path, arcname=image_path.name)
        return zip_path

    def _build_task_name(
        self, config: TemplateConfig, entries: list[ReportEntry]
    ) -> str:
        if config.summary_field == "signature":
            values = [
                entry.signature.strip()
                for entry in entries
                if entry.signature and entry.signature.strip()
            ]
        else:
            values = [
                entry.companyName.strip()
                for entry in entries
                if entry.companyName.strip()
            ]
        summary = ",".join(values) if values else "未知"
        return f"{config.task_prefix}_{summary}"

    def _normalize_value(self, value: object, multiline: bool) -> str:
        if value is None:
            return ""
        if isinstance(value, list):
            cleaned = [str(item).strip() for item in value if str(item).strip()]
            return "\n".join(cleaned) if multiline else "、".join(cleaned)
        return str(value).strip()

    def _normalize_joined_list(self, value: object, separator: str) -> str:
        if value is None:
            return ""
        if isinstance(value, list):
            return separator.join(
                str(item).strip() for item in value if str(item).strip()
            )
        text = str(value).strip()
        if not text:
            return ""
        parts = [
            segment.strip()
            for segment in re.split(r"[\r\n,，、;；]+", text)
            if segment.strip()
        ]
        return separator.join(parts)

    def _safe_name(self, value: str) -> str:
        safe = "".join(char for char in value if char not in '<>:"/\\|?*').strip()
        return safe or uuid4().hex[:8]

    def _text_width(
        self, text: str, font: ImageFont.FreeTypeFont | ImageFont.ImageFont
    ) -> int:
        bbox = font.getbbox(text)
        return bbox[2] - bbox[0]

    def _load_font(self, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
        if self.font_path is not None:
            return ImageFont.truetype(str(self.font_path), size=size)
        return ImageFont.load_default()

    def _resolve_font_path(self) -> Path | None:
        for candidate in DEFAULT_FONT_CANDIDATES:
            path = Path(candidate)
            if path.exists():
                return path
        return None


def qn(tag: str) -> str:
    from docx.oxml.ns import qn as docx_qn

    return docx_qn(tag)
